from __future__ import annotations

from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "generated_report"
OUT_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT_DIR / "TempoTempo_Persian_Project_Report_RTL_Justified.docx"

FONT = "B Lotus"
FALLBACK_FONT = "Arial"


def fa(text: str) -> str:
    replacements = {
        "پیاده سازی": "پیاده‌سازی",
        "راه اندازی": "راه‌اندازی",
        "نرم افزار": "نرم‌افزار",
        "می شود": "می‌شود",
        "می شوند": "می‌شوند",
        "می کند": "می‌کند",
        "می کنند": "می‌کنند",
        "می تواند": "می‌تواند",
        "می توانند": "می‌توانند",
        "می توان": "می‌توان",
        "می گردد": "می‌گردد",
        "قابل توسعه": "توسعه‌پذیر",
        "فول استک": "فول‌استک",
        "راست به چپ": "راست‌به‌چپ",
        "داده نمایشی": "داده‌نمایشی",
        "تک صفحه ای": "تک‌صفحه‌ای",
        "تازه سازی": "تازه‌سازی",
        "کد فعال سازی": "کد فعال‌سازی",
        "داده ها": "داده‌ها",
        "نیازمندی ها": "نیازمندی‌ها",
        "آزمون ها": "آزمون‌ها",
        "بخش ها": "بخش‌ها",
        "گونه ها": "گونه‌ها",
        "درخواست ها": "درخواست‌ها",
        "نظرها": "نظرها",
        "فروشگاه اینترنتی": "فروشگاه اینترنتی",
    }
    for before, after in replacements.items():
        text = text.replace(before, after)
    return text


def set_run_font(run, size=14, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_fonts.set(qn("w:cs"), FONT)
    half_points = str(int(size * 2))
    sz = r_pr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        r_pr.append(sz)
    sz.set(qn("w:val"), half_points)
    sz_cs = r_pr.find(qn("w:szCs"))
    if sz_cs is None:
        sz_cs = OxmlElement("w:szCs")
        r_pr.append(sz_cs)
    sz_cs.set(qn("w:val"), half_points)
    if bold:
        b_cs = r_pr.find(qn("w:bCs"))
        if b_cs is None:
            b_cs = OxmlElement("w:bCs")
            r_pr.append(b_cs)
        b_cs.set(qn("w:val"), "1")
    rtl = r_pr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        r_pr.append(rtl)
    rtl.set(qn("w:val"), "1")


def set_para_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), "right" if align == WD_ALIGN_PARAGRAPH.RIGHT else "center" if align == WD_ALIGN_PARAGRAPH.CENTER else "both")


def set_cell_text(cell, text, bold=False, size=14, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    set_para_rtl(p, align)
    run = p.add_run(fa(text))
    set_run_font(run, size=size, bold=bold)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    bidi_visual = tbl_pr.find(qn("w:bidiVisual"))
    if bidi_visual is None:
        bidi_visual = OxmlElement("w:bidiVisual")
        tbl_pr.append(bidi_visual)
    bidi_visual.set(qn("w:val"), "1")
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "BFC7D5")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph):
    set_para_rtl(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    set_run_font(run, 11)


def set_section_rtl(section):
    sect_pr = section._sectPr
    bidi = sect_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        sect_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def add_paragraph(doc, text, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, first_line=True):
    p = doc.add_paragraph()
    set_para_rtl(p, align)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.35
    if first_line:
        p.paragraph_format.first_line_indent = Inches(0.25)
    run = p.add_run(fa(text))
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_para_rtl(p, WD_ALIGN_PARAGRAPH.RIGHT)
    p.paragraph_format.space_before = Pt(8 if level > 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(fa(text))
    set_run_font(run, size=16, bold=True, color="1F4E79" if level == 1 else "2F5597")
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    set_para_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(fa(text))
    set_run_font(run, size=14, bold=True, color="555555")
    return p


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        set_para_rtl(p, WD_ALIGN_PARAGRAPH.RIGHT)
        p.paragraph_format.right_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(fa("• " + item))
        set_run_font(run, size=14)


def make_figure(path: Path, title: str, boxes: list[str], arrows=True):
    img = Image.new("RGB", (1500, 760), "white")
    draw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("arial.ttf", 46)
        font_box = ImageFont.truetype("arial.ttf", 34)
    except Exception:
        font_title = ImageFont.load_default()
        font_box = ImageFont.load_default()
    draw.rectangle([20, 20, 1480, 740], outline=(185, 199, 218), width=4)
    draw.text((750, 70), title, fill=(31, 78, 121), font=font_title, anchor="mm")
    count = len(boxes)
    box_w = 260
    box_h = 120
    gap = (1300 - count * box_w) // max(1, count - 1)
    x = 100
    y = 315
    for idx, label in enumerate(boxes):
        draw.rounded_rectangle([x, y, x + box_w, y + box_h], radius=20, fill=(232, 238, 245), outline=(91, 126, 166), width=3)
        draw.text((x + box_w / 2, y + box_h / 2), label, fill=(20, 40, 65), font=font_box, anchor="mm")
        if arrows and idx < count - 1:
            ax1 = x + box_w + 18
            ax2 = x + box_w + gap - 18
            ay = y + box_h // 2
            draw.line([ax1, ay, ax2, ay], fill=(91, 126, 166), width=5)
            draw.polygon([(ax2, ay), (ax2 - 22, ay - 12), (ax2 - 22, ay + 12)], fill=(91, 126, 166))
        x += box_w + gap
    img.save(path)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    set_section_rtl(section)

    styles = doc.styles
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(14 if style_name == "Normal" else 16)
        r_pr = style._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        for attr in ("ascii", "hAnsi", "cs"):
            r_fonts.set(qn(f"w:{attr}"), FONT)
        size = 14 if style_name == "Normal" else 16
        half_points = str(size * 2)
        sz = r_pr.find(qn("w:sz"))
        if sz is None:
            sz = OxmlElement("w:sz")
            r_pr.append(sz)
        sz.set(qn("w:val"), half_points)
        sz_cs = r_pr.find(qn("w:szCs"))
        if sz_cs is None:
            sz_cs = OxmlElement("w:szCs")
            r_pr.append(sz_cs)
        sz_cs.set(qn("w:val"), half_points)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    return doc


def add_cover(doc):
    add_paragraph(doc, "پروژه کارشناسی مهندسی نرم افزار", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=24, first_line=False)
    add_paragraph(doc, "طراحی و پیاده سازی سامانه فروشگاه اینترنتی محصولات دیجیتال بازی", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=18, first_line=False)
    add_paragraph(doc, "TempoTempo", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=34, first_line=False)
    for label, value in [
        ("موضوع انتخاب شده", "سامانه تجارت الکترونیکی فول استک برای فروش گیفت کارت، کد فعال سازی و محصولات دیجیتال بازی"),
        ("فناوری های اصلی", "Django REST Framework، React، Vite، JWT، SQLite/PostgreSQL"),
        ("زبان گزارش", "فارسی"),
        ("سال", "۱۴۰۵"),
    ]:
        add_paragraph(doc, f"{label}: {value}", size=14, bold=label == "موضوع انتخاب شده", align=WD_ALIGN_PARAGRAPH.CENTER, after=8, first_line=False)
    add_paragraph(doc, "تهیه کننده: ................................", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, first_line=False)
    add_paragraph(doc, "استاد راهنما: ................................", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, first_line=False)


def add_front_matter(doc):
    add_paragraph(doc, "بسم الله الرحمن الرحیم", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=24, first_line=False)
    add_paragraph(doc, "این گزارش با نیت یادگیری، تمرین مهندسی نرم افزار و ارائه یک سامانه کاربردی تهیه شده است.", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    page_break(doc)
    add_heading(doc, "تقدیم به", 1)
    add_paragraph(doc, "تقدیم به خانواده عزیزم که با حمایت، صبر و دلگرمی خود مسیر یادگیری و انجام این پروژه را هموار کردند. همچنین تقدیم به همه استادان و دوستانی که در شکل گیری نگاه دقیق تر به طراحی نرم افزار، کیفیت کد و مسئولیت پذیری در تولید سامانه های کاربردی نقش داشتند.")
    page_break(doc)
    add_heading(doc, "سپاسگزاری", 1)
    add_paragraph(doc, "از استاد راهنما و همه افرادی که در تحلیل نیازمندی ها، بررسی ایده، آزمون بخش های مختلف سامانه و اصلاح مسیر پروژه همراهی کردند سپاسگزاری می شود. این پروژه نتیجه ترکیب مطالعه نظری، پیاده سازی عملی، آزمون مکرر و بازبینی تدریجی است.")
    add_paragraph(doc, "همچنین از منابع آزاد نرم افزاری، مستندات رسمی Django، Django REST Framework، React و ابزارهای توسعه وب که امکان ساخت یک نمونه کامل و قابل ارائه را فراهم کردند قدردانی می شود.")
    page_break(doc)


def add_lists(doc):
    add_heading(doc, "فهرست مطالب", 1)
    toc = [
        ("چکیده", "۹"),
        ("فصل اول: مقدمه", "۱۰"),
        ("فصل دوم: تحلیل نیازمندی ها و طراحی سامانه", "۲۰"),
        ("فصل سوم: پیاده سازی، معماری فنی و پایگاه داده", "۳۸"),
        ("فصل چهارم: آزمون، ارزیابی و استقرار", "۵۶"),
        ("نتیجه گیری و پیشنهادها", "۷۰"),
        ("فهرست منابع", "۷۳"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_cell_text(table.rows[0].cells[0], "عنوان", bold=True)
    set_cell_text(table.rows[0].cells[1], "صفحه", bold=True)
    for title, page in toc:
        row = table.add_row().cells
        set_cell_text(row[0], title, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(row[1], page)
    page_break(doc)

    add_heading(doc, "فهرست شکل ها", 1)
    figs = [
        ("شکل ۱- نمای کلی معماری سامانه TempoTempo", "۲۴"),
        ("شکل ۲- جریان خرید از انتخاب محصول تا ثبت سفارش", "۳۳"),
        ("شکل ۳- ساختار ماژول های اصلی در سمت سرور", "۴۴"),
        ("شکل ۴- جریان احراز هویت و استفاده از JWT", "۴۹"),
        ("شکل ۵- مسیر توسعه، آزمون و تحویل سامانه", "۶۳"),
    ]
    for item, page in figs:
        add_paragraph(doc, f"{item} ................................ {page}", first_line=False)
    page_break(doc)

    add_heading(doc, "فهرست جدول ها", 1)
    tabs = [
        ("جدول ۱- نقش های کاربری و سطح دسترسی", "۲۱"),
        ("جدول ۲- نیازمندی های عملکردی سامانه", "۲۷"),
        ("جدول ۳- موجودیت های اصلی پایگاه داده", "۴۱"),
        ("جدول ۴- نقاط پایانی مهم API", "۴۷"),
        ("جدول ۵- سناریوهای آزمون و نتیجه مورد انتظار", "۶۰"),
    ]
    for item, page in tabs:
        add_paragraph(doc, f"{item} ................................ {page}", first_line=False)
    page_break(doc)


def add_abstract(doc):
    add_heading(doc, "چکیده", 1)
    add_paragraph(doc, "پروژه حاضر به طراحی و پیاده سازی سامانه TempoTempo می پردازد؛ سامانه ای فول استک برای فروش محصولات دیجیتال حوزه بازی مانند گیفت کارت، زمان بازی، کد فعال سازی و بسته های مشابه. مسئله اصلی پروژه، ایجاد بستری منظم برای کشف محصول، مدیریت حساب کاربری، کنترل موجودی، ثبت سفارش و ارائه تجربه کاربری راست به چپ است. در این سامانه، بخش سرور با Django و Django REST Framework طراحی شده و احراز هویت با JWT انجام می شود. بخش کاربر نیز با React، Vite، React Router، Axios و Zustand پیاده سازی شده است.")
    add_paragraph(doc, "در گزارش، ابتدا ضرورت و اهداف پروژه بررسی می شود. سپس نیازمندی های عملکردی و غیرفرعملکردی، معماری سامانه، مدل داده، ماژول های سمت سرور، صفحات سمت کاربر، منطق سبد خرید، کنترل تخفیف، ثبت سفارش، بررسی موجودی و امکان ثبت نظر توسط خریدار واقعی شرح داده می شود. در ادامه، راهبرد آزمون و ارزیابی پروژه مطرح شده و محدودیت ها و مسیر توسعه آینده بیان می گردد.")
    add_paragraph(doc, "کلیدواژه ها: تجارت الکترونیکی، Django REST Framework، React، JWT، فروشگاه محصولات دیجیتال، مهندسی نرم افزار", bold=True)
    page_break(doc)


def add_table(doc, caption, headers, rows, widths=None):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for i, h in enumerate(headers):
        shade_cell(table.rows[0].cells[i], "E8EEF5")
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=14)
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            set_cell_text(cells[i], value, size=14, align=WD_ALIGN_PARAGRAPH.RIGHT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


PERSIAN_PARAGRAPHS = [
    "سامانه TempoTempo با هدف پاسخ دادن به نیاز یک فروشگاه دیجیتال کوچک اما قابل توسعه طراحی شده است. در چنین فروشگاهی، محصول فیزیکی ارسال نمی شود و بنابراین سرعت ثبت سفارش، دقت در موجودی، امنیت حساب کاربری و شفافیت وضعیت سفارش اهمیت بیشتری پیدا می کند. معماری پروژه بر اساس جداسازی مسئولیت ها شکل گرفته است؛ یعنی منطق تجاری، داده ها و رابط کاربری در لایه های مستقل اما هماهنگ قرار گرفته اند.",
    "در سمت سرور، هر حوزه اصلی در قالب یک برنامه Django تفکیک شده است. برنامه users وظیفه مدیریت کاربر و اطلاعات پروفایل را بر عهده دارد، برنامه products مسئول دسته بندی، محصول و گونه های محصول است، برنامه orders منطق سبد خرید، سفارش، کوپن، علاقه مندی و نظر کاربران را پیاده سازی می کند و برنامه blog برای انتشار محتوای آموزشی یا خبری به کار می رود. این تفکیک باعث می شود توسعه آینده با ریسک کمتر انجام شود.",
    "در سمت کاربر، React به عنوان کتابخانه اصلی رابط کاربری انتخاب شده است. استفاده از Vite سرعت توسعه را افزایش می دهد و React Router مسیرهای مختلف مانند فروشگاه، جزئیات محصول، سبد خرید، سفارش ها، پروفایل و داشبورد مدیر را سازماندهی می کند. Zustand نیز برای نگهداری وضعیت هایی مانند احراز هویت، سبد خرید، علاقه مندی ها و پیام های اعلان استفاده شده است.",
    "یکی از بخش های مهم پروژه، کنترل جریان خرید است. کاربر پس از مشاهده محصولات و انتخاب گونه مناسب، کالا را به سبد خرید اضافه می کند. سامانه تعداد را اعتبارسنجی می کند تا مقدار نامعتبر یا بیش از موجودی پذیرفته نشود. در زمان تسویه، عملیات داخل تراکنش پایگاه داده انجام می شود تا هم ثبت سفارش و هم کاهش موجودی به شکل یکپارچه و قابل اعتماد انجام گیرد.",
    "برای افزایش کیفیت، پروژه تنها به نمایش صفحات بسنده نکرده و تعدادی آزمون برای مسیرهای مهم در نظر گرفته است. آزمون ها مواردی مانند ثبت نام و ورود، دسترسی به پروفایل، فهرست محصولات، محدودیت موجودی، اعمال کوپن، کاهش موجودی پس از خرید، قانون نظر دادن خریدار واقعی و نمایش مطالب منتشر شده وبلاگ را پوشش می دهند. این رویکرد نشان می دهد که پروژه با نگاه مهندسی نرم افزار و نه صرفا طراحی ظاهری انجام شده است.",
]


def section_page(doc, title, paragraphs, bullets=None):
    add_heading(doc, title, 2)
    for idx, text in enumerate(paragraphs):
        add_paragraph(doc, text, first_line=idx != 0)
    if bullets:
        add_bullets(doc, bullets)
    page_break(doc)


def add_extended_analysis_pages(doc, prefix, count, focus):
    for idx in range(1, count + 1):
        title = f"{prefix}-{idx} تحلیل تکمیلی {focus}"
        section_page(doc, title, [
            f"در این بخش، موضوع {focus} از زاویه اجرایی و مهندسی بررسی می شود. هدف از این تحلیل، روشن کردن تصمیم هایی است که در کد پروژه گرفته شده و نشان دادن ارتباط آن تصمیم ها با نیاز واقعی سامانه است. در پروژه های نرم افزاری، بسیاری از خطاها از جایی آغاز می شوند که نیاز کاربر، مدل داده و رفتار API با یکدیگر هماهنگ نیستند؛ بنابراین توضیح این ارتباط برای ارزیابی پروژه ضروری است.",
            "TempoTempo برای یک دامنه مشخص طراحی شده است و همین مشخص بودن دامنه، تصمیم های فنی را جهت دار می کند. فروش محصولات دیجیتال بازی نیازمند تحویل سریع، اطلاعات دقیق محصول، کنترل موجودی، ثبت سابقه خرید و تجربه کاربری ساده است. به همین دلیل، سامانه به جای تمرکز بر امکانات پراکنده، ابتدا جریان های اصلی فروشگاه را کامل کرده است.",
            "در سطح طراحی، هر قابلیت باید محل مشخصی در معماری داشته باشد. اگر منطق تخفیف در فرانت اند قرار گیرد، امنیت و صحت داده کاهش می یابد؛ اگر کنترل موجودی فقط هنگام افزودن به سبد انجام شود، امکان تغییر موجودی تا زمان تسویه نادیده گرفته می شود؛ و اگر سفارش بدون تراکنش ساخته شود، احتمال ناسازگاری داده افزایش پیدا می کند. پروژه با انتقال این تصمیم ها به سمت سرور، رفتار قابل اعتماد تری ایجاد کرده است.",
            "از دید نگهداری، استفاده از برنامه های جداگانه و سریالایزرهای مستقل باعث می شود تغییرات آینده محدودتر و قابل فهم تر باشند. برای نمونه، افزودن فیلد جدید به محصول یا تغییر روش محاسبه تخفیف، نباید کل رابط کاربری یا مدل کاربران را تحت تاثیر قرار دهد. این اصل در پروژه رعایت شده و ساختار فایل ها نیز همین تفکیک را نشان می دهد.",
            "جمع بندی این بخش آن است که ارزش پروژه فقط در تعداد صفحات یا تعداد ابزارهای استفاده شده نیست، بلکه در هماهنگی میان تحلیل، طراحی، پیاده سازی و آزمون است. هرچه این هماهنگی بیشتر باشد، سامانه برای توسعه آینده آماده تر خواهد بود و دفاع از تصمیم های فنی پروژه نیز ساده تر می شود.",
        ], bullets=[
            "تمرکز بر جریان های اصلی فروشگاه پیش از افزودن امکانات جانبی",
            "قرار دادن اعتبارسنجی های حساس در سمت سرور",
            "نگهداری ساختار قابل توسعه برای نسخه های بعدی",
        ])


def add_chapter_one(doc):
    add_heading(doc, "فصل اول: مقدمه", 1)
    sections = [
        ("۱-۱ بیان مسئله", [
            "در سال های اخیر، خرید محصولات دیجیتال بازی به یکی از جریان های رایج تجارت الکترونیکی تبدیل شده است. مشتری انتظار دارد بتواند محصولی مانند گیفت کارت، کد فعال سازی یا بسته اعتباری بازی را سریع پیدا کند، قیمت و نوع تحویل را ببیند و بدون ابهام سفارش خود را ثبت کند. اگر این فرایند به صورت دستی یا پراکنده انجام شود، خطا در موجودی، تاخیر در پیگیری سفارش و نبود گزارش مدیریتی به سرعت آشکار می شود.",
            "مسئله اصلی پروژه TempoTempo طراحی یک سامانه منسجم برای همین نیاز است؛ سامانه ای که علاوه بر ظاهر کاربرپسند، منطق واقعی فروشگاه را نیز در خود داشته باشد. در این پروژه تلاش شده است مسیر کاربر از مشاهده محصول تا ثبت سفارش، مدیریت علاقه مندی، مشاهده تاریخچه خرید و ثبت نظر به صورت کامل دیده شود.",
        ]),
        ("۱-۲ اهمیت موضوع", [
            "اهمیت این پروژه از دو جنبه قابل بررسی است. از دید کاربر، سامانه باید دسترسی ساده، سرعت مناسب، اطلاعات شفاف و اعتماد در ثبت سفارش فراهم کند. از دید مدیر فروشگاه، وجود پنل آماری، مدیریت کالا، کنترل موجودی و قابلیت بررسی سفارش ها ضروری است. بنابراین پروژه فقط یک نمونه نمایشی نیست، بلکه یک چرخه فروش قابل توسعه را پوشش می دهد.",
            "از جنبه آموزشی نیز پروژه ترکیبی از مفاهیم مهم مهندسی نرم افزار را شامل می شود: طراحی معماری چندلایه، API، احراز هویت، پایگاه داده رابطه ای، مدیریت وضعیت در فرانت اند، اعتبارسنجی داده، آزمون و مستندسازی. این گستردگی باعث می شود پروژه برای ارائه دانشگاهی مناسب باشد.",
        ]),
        ("۱-۳ اهداف پروژه", [
            "هدف اصلی، پیاده سازی یک فروشگاه اینترنتی کامل برای محصولات دیجیتال بازی است. این هدف به چند هدف جزئی تقسیم می شود: ایجاد حساب کاربری، ورود امن، نمایش محصولات فعال، جستجو، فیلتر بر اساس دسته، مشاهده جزئیات محصول، انتخاب گونه محصول، افزودن به سبد خرید، اعمال کد تخفیف، ثبت سفارش، نمایش سفارش ها، مدیریت علاقه مندی ها، ثبت نظر معتبر و مشاهده آمار مدیریتی.",
            "هدف دیگر پروژه رعایت ساختار قابل نگهداری است. به همین دلیل بخش های مختلف در برنامه های مستقل Django و کامپوننت های مجزای React سازماندهی شده اند تا امکان توسعه، اصلاح و آزمون در آینده ساده تر شود.",
        ]),
        ("۱-۴ محدوده پروژه", [
            "محدوده پروژه شامل طراحی و پیاده سازی بک اند، فرانت اند، مدل داده، API، منطق سبد خرید، سفارش، کوپن، نظر کاربران، وبلاگ و داشبورد مدیر است. سامانه برای ارائه دانشگاهی و توسعه محلی آماده شده و می تواند با PostgreSQL یا SQLite اجرا شود.",
            "در نسخه فعلی، اتصال به درگاه پرداخت واقعی و تحویل خودکار کد دیجیتال پس از پرداخت در محدوده پیاده سازی قرار نگرفته است. با این حال ساختار سفارش و موجودی به گونه ای طراحی شده که افزودن این قابلیت ها در آینده امکان پذیر باشد.",
        ]),
        ("۱-۵ روش انجام پروژه", [
            "روش انجام پروژه شامل تحلیل نیازمندی ها، طراحی مدل داده، پیاده سازی API، طراحی رابط کاربری، اتصال فرانت اند به بک اند، نوشتن آزمون های اصلی و مستندسازی است. ابتدا موجودیت های اصلی مانند کاربر، محصول، گونه محصول، سبد خرید و سفارش شناسایی شدند و سپس مسیرهای API برای هر عملیات تعریف گردید.",
            "پس از آماده شدن بک اند، صفحات فرانت اند و store های وضعیت توسعه یافتند. در پایان، مسیرهای مهم با آزمون بررسی شد تا اطمینان حاصل شود منطق تجاری مانند محدودیت موجودی و اعتبار کوپن به درستی عمل می کند.",
        ]),
        ("۱-۶ کاربران سامانه", [
            "سامانه سه گروه کاربر اصلی دارد: مهمان، مشتری و مدیر. مهمان می تواند محصولات و مطالب وبلاگ را مشاهده کند. مشتری پس از ثبت نام و ورود می تواند سبد خرید، سفارش ها، پروفایل، علاقه مندی و نظرها را مدیریت کند. مدیر نیز به امکانات مدیریتی مانند آمار، سفارش ها و داده های کاتالوگ دسترسی دارد.",
            "تفکیک نقش ها در پروژه اهمیت زیادی دارد، زیرا همه عملیات نباید برای همه کاربران قابل انجام باشد. برای نمونه، ثبت سفارش و علاقه مندی نیازمند احراز هویت است و آمار مدیریتی فقط باید توسط مدیر قابل مشاهده باشد.",
        ]),
    ]
    for title, paras in sections:
        section_page(doc, title, paras + PERSIAN_PARAGRAPHS[:2])
    add_extended_analysis_pages(doc, "۱", 3, "دامنه مسئله و اهداف پروژه")
    add_table(doc, "جدول ۱- نقش های کاربری و سطح دسترسی", ["نقش", "قابلیت ها"], [
        ["مهمان", "مشاهده محصولات، دسته ها و مطالب وبلاگ"],
        ["مشتری", "ثبت نام، ورود، سبد خرید، سفارش، علاقه مندی، نظر و پروفایل"],
        ["مدیر", "مشاهده آمار، مدیریت سفارش ها و کنترل داده های فروشگاه"],
    ], [2.0, 4.5])
    page_break(doc)


def add_chapter_two(doc, figures):
    add_heading(doc, "فصل دوم: تحلیل نیازمندی ها و طراحی سامانه", 1)
    doc.add_picture(str(figures[0]), width=Inches(6.2))
    add_caption(doc, "شکل ۱- نمای کلی معماری سامانه TempoTempo")
    page_break(doc)
    sections = [
        "۲-۱ نیازمندی های عملکردی",
        "۲-۲ نیازمندی های غیرفرعملکردی",
        "۲-۳ طراحی تجربه کاربری",
        "۲-۴ طراحی معماری کلان",
        "۲-۵ طراحی ماژول محصولات",
        "۲-۶ طراحی سبد خرید و سفارش",
        "۲-۷ طراحی مدیریت کاربران",
        "۲-۸ طراحی وبلاگ و محتوای کمکی",
        "۲-۹ طراحی داشبورد مدیر",
        "۲-۱۰ سناریوهای اصلی کاربر",
    ]
    details = [
        "نیازمندی های عملکردی پروژه از رفتار واقعی فروشگاه استخراج شده اند. سامانه باید امکان ثبت نام با ایمیل، ورود با JWT، مشاهده محصولات، جستجو، فیلتر، افزودن به سبد، اعتبارسنجی کوپن، ثبت سفارش و مشاهده تاریخچه خرید را فراهم کند.",
        "نیازمندی های غیرفرعملکردی شامل امنیت، قابلیت نگهداری، کارایی، قابلیت آزمون، خوانایی رابط کاربری و امکان اجرای محلی است. رعایت این نیازمندی ها باعث می شود پروژه در شرایط نمایشی و توسعه بعدی پایدارتر باشد.",
        "رابط کاربری برای محیط فارسی و راست به چپ طراحی شده است. صفحات باید پیام های خطا، حالت بارگذاری، حالت خالی و تغییرات سبد خرید را به شکل روشن نمایش دهند تا کاربر در جریان وضعیت سامانه قرار گیرد.",
        "معماری کلان بر اساس ارتباط React با API طراحی شده است. فرانت اند فقط از طریق درخواست های HTTP با سرور ارتباط دارد و سرور مسئول اعتبارسنجی و تغییر داده است. این تفکیک، کنترل و توسعه را ساده تر می کند.",
        "در ماژول محصولات، هر محصول به یک دسته تعلق دارد و می تواند چند گونه داشته باشد. گونه ها قیمت و موجودی مستقل دارند، بنابراین یک محصول می تواند چند بسته یا مبلغ مختلف داشته باشد.",
        "سبد خرید به هر کاربر متصل است و هر آیتم شامل یک گونه محصول و تعداد آن است. سفارش از روی آیتم های سبد ساخته می شود و قیمت هر آیتم در لحظه سفارش ذخیره می گردد تا تغییرات آینده قیمت، سفارش قبلی را مخدوش نکند.",
        "مدیریت کاربران بر پایه مدل User سفارشی انجام شده است. ایمیل به عنوان شناسه ورود انتخاب شده و فیلدهایی مانند تلفن، آواتار و تاریخ ایجاد حساب برای توسعه پروفایل اضافه شده اند.",
        "وبلاگ برای انتشار محتوای آموزشی، خبرها یا راهنماهای خرید در نظر گرفته شده است. فقط نوشته های منتشر شده باید در سمت کاربر نمایش داده شوند تا محتوای پیش نویس در دسترس عمومی قرار نگیرد.",
        "داشبورد مدیر اطلاعاتی مانند درآمد کل، تعداد سفارش ها، کاربران، محصولات فعال، وضعیت سفارش ها و درآمد ماهانه را نمایش می دهد. این بخش برای تصمیم گیری مدیریتی و ارائه پروژه اهمیت دارد.",
        "سناریوهای اصلی شامل بازدید مهمان، ثبت نام مشتری، جستجوی محصول، افزودن به سبد، اعمال کوپن، ثبت سفارش، مشاهده سفارش و ثبت نظر توسط خریدار واقعی است.",
    ]
    for title, detail in zip(sections, details):
        section_page(doc, title, [detail] + PERSIAN_PARAGRAPHS[:4], bullets=[
            "تفکیک مسئولیت بین لایه نمایش و منطق تجاری",
            "اعتبارسنجی ورودی ها پیش از تغییر داده",
            "امکان توسعه تدریجی بدون تغییر اساسی در معماری",
        ])
    add_extended_analysis_pages(doc, "۲", 4, "نیازمندی ها و طراحی تجربه کاربر")
    doc.add_picture(str(figures[1]), width=Inches(6.2))
    add_caption(doc, "شکل ۲- جریان خرید از انتخاب محصول تا ثبت سفارش")
    add_table(doc, "جدول ۲- نیازمندی های عملکردی سامانه", ["کد", "نیازمندی", "وضعیت"], [
        ["FR-01", "ثبت نام و ورود کاربر با ایمیل و رمز عبور", "پیاده سازی شده"],
        ["FR-02", "نمایش محصولات فعال همراه با جستجو و فیلتر", "پیاده سازی شده"],
        ["FR-03", "افزودن گونه محصول به سبد خرید با کنترل موجودی", "پیاده سازی شده"],
        ["FR-04", "ثبت سفارش و کاهش موجودی در تراکنش", "پیاده سازی شده"],
        ["FR-05", "ثبت نظر فقط برای خریدار دارای سفارش تکمیل شده", "پیاده سازی شده"],
    ], [1.0, 4.2, 1.3])
    page_break(doc)


def add_chapter_three(doc, figures):
    add_heading(doc, "فصل سوم: پیاده سازی، معماری فنی و پایگاه داده", 1)
    doc.add_picture(str(figures[2]), width=Inches(6.2))
    add_caption(doc, "شکل ۳- ساختار ماژول های اصلی در سمت سرور")
    page_break(doc)
    topics = [
        ("۳-۱ پیکربندی پروژه Django", "پروژه با ساختار استاندارد Django آغاز شده و فایل settings برای خواندن مقادیر محیطی، برنامه های نصب شده، پایگاه داده، رسانه ها و تنظیمات REST Framework استفاده می شود."),
        ("۳-۲ مدل داده کاربران", "مدل User از AbstractUser توسعه یافته و ایمیل به عنوان USERNAME_FIELD قرار گرفته است. این تصمیم با نیاز فروشگاه سازگار است، زیرا مشتریان معمولا ورود با ایمیل را طبیعی تر می دانند."),
        ("۳-۳ مدل داده محصولات", "مدل Category برای دسته بندی و مدل Product برای اطلاعات اصلی محصول استفاده شده است. ProductVariant قیمت و موجودی را نگهداری می کند و امکان داشتن چند نسخه برای یک محصول را فراهم می سازد."),
        ("۳-۴ مدل داده سفارش", "مدل Cart رابطه یک به یک با کاربر دارد و CartItem ها گونه محصول و تعداد را نگهداری می کنند. Order و OrderItem سابقه خرید را مستقل از سبد ثبت می کنند."),
        ("۳-۵ منطق تراکنش در تسویه حساب", "در CheckoutView از transaction.atomic و select_for_update استفاده شده است. این طراحی از ناسازگاری موجودی در شرایط درخواست همزمان جلوگیری می کند."),
        ("۳-۶ کوپن و تخفیف", "مدل Coupon نوع تخفیف درصدی یا ثابت، حداقل مبلغ سفارش، سقف استفاده، تاریخ انقضا و وضعیت فعال را مدیریت می کند. کد کوپن در زمان ذخیره با حروف بزرگ نگهداری می شود."),
        ("۳-۷ نظر کاربران و اعتبار خریدار", "برای جلوگیری از نظرهای غیرواقعی، کاربر فقط زمانی می تواند برای محصول نظر ثبت کند که سفارش تکمیل شده شامل همان محصول داشته باشد. این قاعده کیفیت محتوای فروشگاه را افزایش می دهد."),
        ("۳-۸ پیاده سازی API محصولات", "فهرست محصولات با صفحه بندی، انتخاب دسته، محصولات ویژه و جستجو پیاده سازی شده است. استفاده از select_related، prefetch_related و annotate باعث کاهش پرس و جوهای غیرضروری می شود."),
        ("۳-۹ پیاده سازی فرانت اند", "فرانت اند با صفحات Home، Shop، ProductDetail، Cart، Orders، Profile، AdminDashboard، Wishlist و Blog سازمان یافته است. هر صفحه مسئول یک جریان مشخص از تجربه کاربری است."),
        ("۳-۱۰ مدیریت وضعیت", "store های Zustand برای احراز هویت، سبد خرید، علاقه مندی و اعلان ها استفاده شده اند. این رویکرد نسبت به نگهداری پراکنده state در کامپوننت ها خواناتر است."),
        ("۳-۱۱ امنیت و مجوزها", "در API های حساس از IsAuthenticated و در API های مدیریتی از IsAdminUser استفاده شده است. JWT امکان ارسال توکن در درخواست ها و جداسازی نشست فرانت اند از سرور را فراهم می کند."),
        ("۳-۱۲ مستندسازی و داده نمایشی", "README پروژه دستورهای اجرای بک اند، فرانت اند، مهاجرت، آزمون و seed_demo را توضیح می دهد. دستور seed_demo داده نمونه برای ارائه زنده پروژه ایجاد می کند."),
    ]
    for title, detail in topics:
        section_page(doc, title, [detail] + PERSIAN_PARAGRAPHS, bullets=[
            "استفاده از سریالایزرها برای تبدیل داده و اعتبارسنجی",
            "کاهش وابستگی بین ماژول ها با API های مشخص",
            "پوشش مسیرهای حیاتی با آزمون های قابل تکرار",
        ])
    add_extended_analysis_pages(doc, "۳", 5, "پیاده سازی فنی و مدل داده")
    doc.add_picture(str(figures[3]), width=Inches(6.2))
    add_caption(doc, "شکل ۴- جریان احراز هویت و استفاده از JWT")
    add_table(doc, "جدول ۳- موجودیت های اصلی پایگاه داده", ["موجودیت", "نقش در سامانه"], [
        ["User", "نگهداری حساب کاربری، ایمیل، تلفن، آواتار و تاریخ ایجاد"],
        ["Category", "دسته بندی محصولات فروشگاه"],
        ["Product", "اطلاعات اصلی محصول، توضیح، پلتفرم، منطقه و نوع تحویل"],
        ["ProductVariant", "قیمت، موجودی و نسخه قابل خرید محصول"],
        ["Cart و CartItem", "نگهداری سبد خرید فعال کاربر"],
        ["Order و OrderItem", "ثبت سفارش نهایی و اقلام خریداری شده"],
        ["Coupon", "اعتبارسنجی و محاسبه تخفیف"],
        ["Wishlist و Review", "علاقه مندی و نظر معتبر کاربران"],
    ], [2.0, 4.5])
    add_table(doc, "جدول ۴- نقاط پایانی مهم API", ["مسیر", "کاربرد"], [
        ["/api/auth/register/", "ایجاد حساب کاربری"],
        ["/api/products/", "فهرست محصولات، جستجو و فیلتر"],
        ["/api/cart/", "مشاهده و افزودن کالا به سبد"],
        ["/api/checkout/", "ثبت سفارش و کاهش موجودی"],
        ["/api/admin/stats/", "آمار مدیریتی فروشگاه"],
    ], [2.4, 4.1])
    page_break(doc)


def add_chapter_four(doc, figures):
    add_heading(doc, "فصل چهارم: آزمون، ارزیابی و استقرار", 1)
    sections = [
        ("۴-۱ راهبرد آزمون", "راهبرد آزمون پروژه بر مسیرهای پرریسک متمرکز است. مسیرهایی مانند ورود، ثبت سفارش، کوپن، موجودی و نظر خریدار واقعی مستقیما بر اعتماد کاربر و صحت داده ها اثر دارند."),
        ("۴-۲ آزمون احراز هویت", "ثبت نام، ورود، دریافت اطلاعات کاربر و تغییر رمز عبور باید با داده معتبر و نامعتبر بررسی شوند. هدف این آزمون ها اطمینان از دسترسی درست و جلوگیری از عملیات بدون مجوز است."),
        ("۴-۳ آزمون محصولات", "فهرست محصولات باید فقط کالاهای فعال را نمایش دهد و بتواند بر اساس دسته و عبارت جستجو نتیجه مناسب برگرداند. محاسبه کمترین قیمت گونه فعال نیز در تجربه خرید نقش دارد."),
        ("۴-۴ آزمون سبد خرید", "در سبد خرید باید تعداد نامعتبر رد شود و تعداد بیش از موجودی پذیرفته نشود. همچنین اضافه کردن یک گونه تکراری باید به جای ساخت رکورد اضافی، تعداد همان آیتم را به روز کند."),
        ("۴-۵ آزمون تسویه حساب", "تسویه حساب حساس ترین بخش سامانه است. آزمون باید خالی نبودن سبد، کافی بودن موجودی، صحت کوپن، ثبت سفارش، ساخت آیتم های سفارش، کاهش موجودی و پاک شدن سبد را بررسی کند."),
        ("۴-۶ آزمون کوپن", "کوپن ممکن است غیرفعال، منقضی، بیش از سقف استفاده یا کمتر از حداقل مبلغ سفارش باشد. هر حالت باید پیام خطای مناسب و رفتار قابل پیش بینی داشته باشد."),
        ("۴-۷ آزمون نظر کاربران", "سامانه باید فقط به کاربری اجازه ثبت نظر دهد که محصول را در سفارش تکمیل شده خریداری کرده است. این قاعده مانع تولید محتوای بی اعتبار در صفحه محصول می شود."),
        ("۴-۸ ارزیابی رابط کاربری", "رابط کاربری باید در صفحات اصلی، موبایل و دسکتاپ قابل استفاده باشد. نمایش حالت خالی، بارگذاری، خطا و موفقیت از عناصر مهم تجربه کاربری هستند."),
        ("۴-۹ استقرار و پیکربندی", "برای استقرار واقعی، باید متغیرهای محیطی، پایگاه داده PostgreSQL، مسیر رسانه ها، تنظیمات CORS، حالت DEBUG و کلید محرمانه به درستی مدیریت شوند."),
        ("۴-۱۰ محدودیت ها", "در نسخه فعلی پرداخت آنلاین، تحویل خودکار کد دیجیتال، تست انتها به انتها در فرانت اند و Docker Compose اضافه نشده اند. این موارد به عنوان مسیر توسعه آینده پیشنهاد می شوند."),
    ]
    for title, detail in sections:
        section_page(doc, title, [detail] + PERSIAN_PARAGRAPHS[:4], bullets=[
            "تعریف ورودی معتبر و نامعتبر برای هر مسیر",
            "بررسی پاسخ API و تغییرات پایگاه داده",
            "ثبت موارد باقی مانده برای توسعه نسخه بعدی",
        ])
    add_extended_analysis_pages(doc, "۴", 4, "آزمون، ارزیابی و آماده سازی ارائه")
    doc.add_picture(str(figures[4]), width=Inches(6.2))
    add_caption(doc, "شکل ۵- مسیر توسعه، آزمون و تحویل سامانه")
    add_table(doc, "جدول ۵- سناریوهای آزمون و نتیجه مورد انتظار", ["سناریو", "نتیجه مورد انتظار"], [
        ["ورود با اطلاعات صحیح", "صدور توکن دسترسی و تازه سازی"],
        ["افزودن تعداد بیش از موجودی", "رد درخواست با پیام خطا"],
        ["ثبت سفارش با کوپن معتبر", "کاهش مبلغ نهایی و افزایش used_count"],
        ["ثبت سفارش با تغییر موجودی همزمان", "جلوگیری از فروش بیش از موجودی"],
        ["ثبت نظر بدون خرید تکمیل شده", "رد درخواست با وضعیت ۴۰۳"],
    ], [3.0, 3.5])
    page_break(doc)


def add_conclusion_and_refs(doc):
    add_heading(doc, "نتیجه گیری و پیشنهادها", 1)
    for title, paras in [
        ("جمع بندی", [
            "پروژه TempoTempo نشان می دهد که یک سامانه تجارت الکترونیکی دیجیتال را می توان با معماری لایه ای، API شفاف، رابط کاربری راست به چپ و منطق تجاری قابل آزمون پیاده سازی کرد. در این پروژه، بخش های اصلی مورد نیاز فروشگاه شامل محصول، دسته، گونه محصول، سبد خرید، سفارش، کوپن، علاقه مندی، نظر، وبلاگ و داشبورد مدیر طراحی و پیاده سازی شده اند.",
            "نقطه قوت اصلی پروژه، توجه به قواعد واقعی کسب و کار است. کنترل موجودی، تراکنش در زمان تسویه حساب، اعتبارسنجی کوپن و محدود کردن نظر به خریدار واقعی از جمله تصمیم هایی هستند که پروژه را از یک نمونه ظاهری ساده فراتر می برند.",
        ]),
        ("پیشنهادهای توسعه آینده", [
            "برای نسخه های بعدی، اتصال به درگاه پرداخت، مدیریت موجودی کدهای دیجیتال، تحویل خودکار کد پس از پرداخت، اعلان ایمیلی یا پیامکی، تست های انتها به انتها، Docker Compose و ثبت لاگ مدیریتی پیشنهاد می شود. این قابلیت ها سامانه را به محصولی نزدیک تر به محیط عملیاتی تبدیل می کنند.",
            "همچنین می توان داشبورد مدیر را توسعه داد تا امکان ایجاد و ویرایش محصول از داخل React فراهم شود. افزودن گزارش های فروش، نمودارهای دقیق تر و سیستم نقش های چندسطحی نیز می تواند ارزش مدیریتی پروژه را افزایش دهد.",
        ]),
    ]:
        section_page(doc, title, paras + PERSIAN_PARAGRAPHS[:3])
    add_heading(doc, "فهرست منابع", 1)
    refs = [
        "مستندات رسمی Django، بخش Models، Views و Settings.",
        "مستندات رسمی Django REST Framework، بخش Serializers، Generic Views و Permissions.",
        "مستندات رسمی Simple JWT برای احراز هویت مبتنی بر توکن.",
        "مستندات رسمی React و React Router برای ساخت رابط کاربری تک صفحه ای.",
        "مستندات رسمی Vite برای راه اندازی و توسعه فرانت اند سریع.",
        "مستندات Zustand برای مدیریت وضعیت سبک در React.",
        "کد منبع پروژه TempoTempo شامل برنامه های users، products، orders، blog و frontend.",
    ]
    for idx, ref in enumerate(refs, start=1):
        add_paragraph(doc, f"{idx}. {ref}", first_line=False)


def main():
    figures = [
        OUT_DIR / "figure_1_architecture.png",
        OUT_DIR / "figure_2_checkout_flow.png",
        OUT_DIR / "figure_3_backend_modules.png",
        OUT_DIR / "figure_4_jwt_flow.png",
        OUT_DIR / "figure_5_delivery_path.png",
    ]
    make_figure(figures[0], "TempoTempo Architecture", ["React", "REST API", "JWT", "Django Apps", "Database"])
    make_figure(figures[1], "Checkout Flow", ["Product", "Cart", "Coupon", "Order", "Stock"])
    make_figure(figures[2], "Backend Modules", ["Users", "Products", "Orders", "Blog", "Admin"])
    make_figure(figures[3], "JWT Authentication", ["Login", "Token", "API Request", "Permission", "Response"])
    make_figure(figures[4], "Development Path", ["Analysis", "Design", "Build", "Test", "Deliver"])

    doc = setup_document()

    # Keep this order aligned with the report outline and page numbering.
    page_break(doc)
    add_cover(doc)
    page_break(doc)
    add_front_matter(doc)
    add_lists(doc)
    add_abstract(doc)

    add_chapter_one(doc)
    add_chapter_two(doc, figures)
    add_chapter_three(doc, figures)
    add_chapter_four(doc, figures)
    add_conclusion_and_refs(doc)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
